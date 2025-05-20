
from flask import Flask, render_template, request, send_file, url_for
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
import os, re, fitz, matplotlib.pyplot as plt
from wordcloud import WordCloud
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = './uploads'
app.config['REPORT_FOLDER'] = './reports'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['REPORT_FOLDER'], exist_ok=True)
os.makedirs("static/images", exist_ok=True)


def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    return "".join([page.get_text() for page in doc])

def generate_wordcloud(text, path):
    wc = WordCloud(width=400, height=200, background_color='white').generate(text)
    wc.to_file(os.path.join("static/images", path))

def generate_snapshot(pdf_path, image_path):
    doc = fitz.open(pdf_path)
    pix = doc[0].get_pixmap()
    pix.save(os.path.join("static/images", image_path))


def calculate_similarity(text1, text2):
    tfidf = TfidfVectorizer(stop_words='english')
    vectors = tfidf.fit_transform([text1, text2])
    return round(float(cosine_similarity(vectors[0:1], vectors[1:2])[0][0] * 100), 2)

def detect_resume_sections(text):
    has_email = re.search(r'[\w\.-]+@[\w\.-]+', text)
    has_phone = re.search(r'\+?\d[\d\s\-()]{7,}\d', text)
    has_skills = 'skills' in text.lower()
    has_work = 'experience' in text.lower() or 'work' in text.lower()
    has_edu = 'education' in text.lower() or 'degree' in text.lower()
    return {
        'Contact': bool(has_email or has_phone),
        'Skills': has_skills,
        'Work History': has_work,
        'Education': has_edu
    }

def extract_missing_keywords(jd_text, resume_text, top_n=10):
    vectorizer = CountVectorizer(stop_words='english', max_features=top_n)
    jd_text = re.sub(r'[^\w\s]', '', jd_text)
    resume_text = re.sub(r'[^\w\s]', '', resume_text)
    jd_words = vectorizer.fit([jd_text]).get_feature_names_out()
    return [word for word in jd_words if word not in resume_text.lower()]

def save_score_donut(score, path):
    fig, ax = plt.subplots(figsize=(4, 4), subplot_kw=dict(aspect="equal"))
    data = [score, 100 - score]
    wedges, _ = ax.pie(data, wedgeprops=dict(width=0.3), startangle=90, colors=["#FFCF01", "#1C1C1C"])
    ax.text(0, 0, f"{score}%", ha="center", va="center", fontsize=18, fontweight="bold")
    plt.savefig(path, bbox_inches='tight')
    plt.close()

def save_section_bar_chart(section_results, path):
    labels = list(section_results.keys())
    scores = [100 if val else 30 for val in section_results.values()]
    fig, ax = plt.subplots()
    ax.bar(labels, scores, color='#FFCF01')
    ax.set_ylim(0, 100)
    ax.set_ylabel('Confidence')
    ax.set_title('Resume Section Coverage')
    plt.tight_layout()
    plt.savefig(path)
    plt.close()

def generate_recommendation(score, sections):
    rec = []
    if score < 60:
        rec.append("ðŸ”§ Resume needs major improvement.")
    elif score < 80:
        rec.append("ðŸ› ï¸ Resume is moderately aligned.")
    else:
        rec.append("âœ… Resume is well-aligned.")
    for sec, present in sections.items():
        if not present:
            rec.append(f"âš ï¸ Add/improve {sec.lower()} section.")
    return " ".join(rec)

    df = pd.DataFrame({
        'Candidate Name': [name],
        'Email': [email],
        'Job Title': [job],
        'Resume Score': [score],
        'Missing Keywords': [", ".join(keywords)],
        **{k: ["âœ”ï¸" if v else "âŒ"] for k, v in section_results.items()}
    })
    df.to_csv(path, index=False)

def generate_report(candidate_name, email, job_title, score, section_results, wordcloud_path, resume_img, jd_img, donut_path, output_path):
    template = Document()
    template.add_paragraph("Candidate Name: " + candidate_name)
    template.add_paragraph("Email: " + email)
    template.add_paragraph("Job Title: " + job_title)
    template.add_paragraph("Resume Score: " + str(score) + "%")
    template.add_paragraph("Section Analysis:")
    for section, result in section_results.items():
        template.add_paragraph(f"âœ”ï¸ {section}" if result else f"âŒ {section}")
    template.add_picture(os.path.join("static/images", wordcloud_path), width=Inches(3.5))
    template.add_picture(os.path.join("static/images", resume_img), width=Inches(2.5))
    template.add_picture(os.path.join("static/images", jd_img), width=Inches(2.5))
    template.add_picture(os.path.join("static/images", donut_path), width=Inches(3.5))
    template.save(output_path)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            # Ensure Form Fields are Provided
            resume = request.files.get('resume')
            jd = request.files.get('jd')
            name = request.form.get('candidate_name')
            email = request.form.get('email')
            job = request.form.get('job_title')

            if not (resume and jd and name and email and job):
                return "Error: All fields are required.", 400

            # Save Uploaded Files Securely
            resume_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(resume.filename))
            jd_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(jd.filename))
            resume.save(resume_path)
            jd.save(jd_path)

            print("âœ… Files saved successfully.")

            # Extract Text from PDF Files
            resume_text = extract_text_from_pdf(resume_path)
            jd_text = extract_text_from_pdf(jd_path)

            print("âœ… Text extracted from PDFs.")

            # Calculate Resume Score
            score = calculate_similarity(resume_text, jd_text)
            section_results = detect_resume_sections(resume_text)
            missing_keywords = extract_missing_keywords(jd_text, resume_text)

            print(f"âœ… Score Calculated: {score}%")
            print(f"âœ… Section Results: {section_results}")

            # Generate Visuals and Report Paths
            wordcloud_path = os.path.join("static", "images", "wordcloud.png")
            resume_snapshot = os.path.join("static", "images", "resume.png")
            jd_snapshot = os.path.join("static", "images", "jd.png")
            donut_path = os.path.join("static", "images", "donut.png")
            section_bar_path = os.path.join("static", "images", "section_bar.png")

            # Generate Visuals
            generate_wordcloud(jd_text, wordcloud_path)
            generate_snapshot(resume_path, resume_snapshot)
            generate_snapshot(jd_path, jd_snapshot)
            save_score_donut(score, donut_path)
            save_section_bar_chart(section_results, section_bar_path)

            print("âœ… Visuals generated successfully.")

            # Generate Recommendations
            recommendation = generate_recommendation(score, section_results)
            print(f"âœ… Recommendation Generated: {recommendation}")

            # Generate Report Path and File
            report_filename = f"{name.replace(' ', '_')}_report.docx"
            report_path = os.path.join(app.config['REPORT_FOLDER'], report_filename)
            generate_report(name, email, job, score, section_results, 
                            wordcloud_path, resume_snapshot, jd_snapshot, donut_path, report_path)

            print(f"âœ… Report generated at {report_path}")

            # Return Rendered Result Page
            return render_template('index.html',
                                   score=score,
                                   section_results=section_results,
                                   missing_keywords=missing_keywords,
                                   recommendation=recommendation,
                                   donut=url_for('static', filename='images/donut.png'),
                                   section_bar=url_for('static', filename='images/section_bar.png'),
                                   wordcloud=url_for('static', filename='images/wordcloud.png'),
                                   resume_snapshot=url_for('static', filename='images/resume.png'),
                                   jd_snapshot=url_for('static', filename='images/jd.png'))

        except Exception as e:
            print(f"âŒ Error in Processing: {str(e)}")
            return f"Internal Server Error: {str(e)}", 500

    # Display the Form if GET request
    return render_template('index.html')

@app.route('/download_report')
def download_report():
    files = os.listdir(app.config['REPORT_FOLDER'])
    latest = max([os.path.join(app.config['REPORT_FOLDER'], f) for f in files], key=os.path.getctime)
    return send_file(latest, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), threaded=True)

